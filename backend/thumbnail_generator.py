from PIL import Image, ImageDraw, ImageFont
from docx import Document
import os

def generate_thumbnail(template_path, thumbnail_path):
    """
    Generates a thumbnail for the document template.

    Args:
        template_path (str): Path to the .docx template.
        thumbnail_path (str): Path to save the generated thumbnail.
    """
    # Load the .docx file
    try:
        doc = Document(template_path)
    except Exception as e:
        raise ValueError(f"Error loading .docx file: {e}")

    # Extract the first few lines of text from the document
    preview_text = ""
    for paragraph in doc.paragraphs[:len(doc.paragraphs)]:  # Adjust to display the desired number of paragraphs
        preview_text += paragraph.text + "\n"
    preview_text = preview_text.strip() if preview_text else "No content available in template."

    # Create a blank image for the thumbnail
    img = Image.new('RGB', (400, 300), color=(255, 255, 255))  # White background
    draw = ImageDraw.Draw(img)

    # Load font for drawing text
    font_size = 16
    try:
        font = ImageFont.truetype("arial.ttf", font_size)  # Adjust path to arial.ttf if needed
    except IOError:
        font = ImageFont.load_default()  # Use default font if arial.ttf is not available

    # Add text to the image
    margin = 10
    current_height = margin
    for line in preview_text.split("\n"):
        if current_height + font_size > img.height - margin:  # Stop if text exceeds the image height
            draw.text((margin, current_height), "...", fill="black", font=font)
            break
        draw.text((margin, current_height), line, fill="black", font=font)
        current_height += font_size + 5  # Add line spacing

    # Save the thumbnail
    os.makedirs(os.path.dirname(thumbnail_path), exist_ok=True)
    img.save(thumbnail_path)
